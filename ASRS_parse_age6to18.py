# %pip install pypdf
# %pip install pdfplumber
# %pip install pandas
# %pip install python-docx

# %pip install pyinflect spacy
# %python -m spacy download en_core_web_sm

# %pip install pdfminer.six==20221105
# %pip install pdfplumber==0.9.0

import pdfplumber
import pandas as pd

import spacy
import pyinflect

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

nlp = spacy.load('en_core_web_sm')

tables = []
firstName = 'Julian'

######################                          Extract the tables from the pdf                           ###########################
#####################################################################################################################################

import pdfplumber
with pdfplumber.open("C:/Users/ljkop/Laura/ASRS/Julian__Gutierrez.pdf") as pdf:
    for page in pdf.pages:
        tables_on_page = page.extract_tables({})

        if tables_on_page:
            for table in tables_on_page:
                if table:
                    tables.append({
                        'page': pdf.pages.index(page) + 1,
                        'data': table
                    })

def get_table(idx_tables):
    if tables:
        temp_table = tables[idx_tables]['data']
        df_temp = pd.DataFrame(temp_table[1:], columns=temp_table[0])  # first row as header
    return df_temp

#need to add df_sr and change df_asr to df_a. Will need to merge 2 tables of 'unusual behaviors'
df_ASRS = get_table(0)


######################                   Clean up the ASRS table for appendix table                       ###########################
#####################################################################################################################################

## remove unnecessary rows
df_ASRS = df_ASRS[df_ASRS['Scale'] != 'DSM-5 Scale']
df_ASRS = df_ASRS[df_ASRS['Scale'] != 'TOTAL SCORE']
df_ASRS = df_ASRS[df_ASRS['Scale'] != 'ASRS SCALES']
df_ASRS = df_ASRS[df_ASRS['Scale'] != 'DSM-5 SCALE']
df_ASRS = df_ASRS[df_ASRS['Scale'] != 'TREATMENT SCALES']

## move total to bottom row
df_ASRS = pd.concat([df_ASRS.iloc[1:], df_ASRS.iloc[[0]]]).reset_index(drop=True)

## remove the \n and everything after in the t-score column
df_ASRS.iloc[:, 1] = df_ASRS.iloc[:, 1].str.split('\n').str[0]

#remove the \n and everything after in the column header
df_ASRS.columns = list(df_ASRS.columns[:1]) + [df_ASRS.columns[1].split('\n')[0]] + list(df_ASRS.columns[2:])

## Remove remaining \n
df_ASRS = df_ASRS.apply(lambda col: col.map(lambda x: x.replace('\n', ' ') if isinstance(x, str) else x))

df_ASRSappendix = df_ASRS.iloc[:, [0, 1, 2, 3]]

doc = Document()

# Add table with header row + data rows
table = doc.add_table(rows=1, cols=len(df_ASRSappendix.columns))
table.style = 'Table Grid'

# Set column widths (in inches converted to EMUs)
from docx.shared import Inches
table.columns[0].width = Inches(2.5)  # first column wide
table.columns[1].width = Inches(0.75)  # middle column narrow
table.columns[2].width = Inches(0.75)  # middle column narrow
table.columns[3].width = Inches(2.5)  # first column wide

# Write header row (bolded)
header_cells = table.rows[0].cells
for i, col_name in enumerate(df_ASRSappendix.columns):
    header_cells[i].text = str(col_name)
    for run in header_cells[i].paragraphs[0].runs:
        run.bold = True

# Write data rows
for row_idx, (_, row) in enumerate(df_ASRSappendix.iterrows()):
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value) if value is not None else ''
    
    # Bold the last row
    if row_idx == len(df_ASRSappendix) - 1:
        for cell in row_cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.save("C:/Users/ljkop/Laura/ASRS/ASRSappendixTable.docx")

######################                   Clean and condition the item tables                              ###########################
#####################################################################################################################################

def clean_table(df_temp, col_name):

    ## replace the column name from 'none' to 'Item Score'
    df_temp.columns = ['Item Score' if col is None else col for col in df_temp.columns]
   
    ## remove the first row as it is a remnant of the orig column name
    df_temp = df_temp.iloc[1:].reset_index(drop=True)

    ## remove any rows that contain raw scores as they don't contain a comment
    df_temp = remove_raw_score_rows(df_temp)
    
    ## remove number at beginning of sentence fragment
    df_temp[col_name] = df_temp[col_name].str.replace(r'^\d+\.\s*', '', regex=True)

    ## remove '?' from the end of each sentence
    df_temp[col_name] = df_temp[col_name].str.rstrip('?')

    ## sort by score highest to lowest
    df_temp = df_temp.sort_values(by='Item Score', ascending=False).reset_index(drop=True)
    
    ##  negate sentnece if reverse scoring
    df_temp[col_name] = df_temp[col_name].apply(negate_sentence)

    ##  make first verb singular tense
    df_temp[col_name] = df_temp[col_name].apply(singularize_first_verb)

    ##  make first verb singular tense
    df_temp[col_name] = df_temp[col_name].apply(to_present_tense)
    
    ##  add first name to sentence
    df_temp[col_name] = firstName + ' ' + df_temp[col_name] + '.'
        
    return df_temp

def singularize_first_verb(text):
    words = text.split()
    if words:
        doc = nlp(words[0])
        inflected = doc[0]._.inflect('VBZ')  # third person singular present
        if inflected:
            words[0] = inflected
    return ' '.join(words)

def to_present_tense(text):
    doc = nlp(text)
    result = []
    for token in doc:
        if token.tag_ in ('VBD', 'VBP', 'VBG', 'VBN'):  # past, non-3rd singular, gerund, past participle
            inflected = token._.inflect('VBZ')  # convert to third person singular present
            result.append(inflected if inflected else token.text)
        else:
            result.append(token.text)
    return ' '.join(result)

def negate_sentence(text):
    if '(R)' not in text:
        return text
    clean = text.split('?')[0].strip()
    return 'does not ' + clean

## remove any row with text = "Raw Score =" as this row does not contain a comment
def remove_raw_score_rows(df_temp):
    # Get the name of the first column
    first_col = df_temp.columns[0]
    
    # Keep only rows where the first column does NOT contain "Raw Score ="
    mask = ~df_temp[first_col].astype(str).str.contains('Raw Score =', na=False)
    print(mask)
    df_cleaned = df_temp[mask].reset_index(drop=True)
    
    return df_cleaned


df_ub_cont = get_table(2)
df_ub_cont = df_ub_cont.rename(columns={'Unusual Behaviors (continued)': 'Unusual Behaviors'})
df_sc = get_table(3)
df_sr = get_table(4)
df_ub = get_table(5)
df_ps = get_table(7)
df_as = get_table(8)
df_ser = get_table(9)
df_al = get_table(10)
df_ss = get_table(11)
df_s = get_table(12)
df_a = get_table(13)
df_br = get_table(14)

#print(df_sc)

df_ub = clean_table(df_ub, 'Unusual Behaviors')
df_ub_cont = clean_table(df_ub_cont, 'Unusual Behaviors')
df_sc = clean_table(df_sc, 'Social/Communication')
df_sr = clean_table(df_sr, 'Self-Regulation')
df_ps = clean_table(df_ps, 'Peer Socialization')
df_as = clean_table(df_as, 'Adult Socialization')
df_ser = clean_table(df_ser, 'Social/Emotional Reciprocity')
df_al = clean_table(df_al, 'Atypical Language')
df_ss = clean_table(df_ss, 'Sensory Sensitivity')
df_s = clean_table(df_s, 'Stereotypy')
df_a = clean_table(df_a, 'Attention')
df_br = clean_table(df_br, 'Behavioral Rigidity')

df_ub = pd.concat([df_ub, df_ub_cont], ignore_index=True)

## sort by score highest to lowest
df_ub = df_ub.sort_values(by='Item Score', ascending=False).reset_index(drop=True)

#####################                   Finish conditioning ASRS table for body                          ###########################
####################################################################################################################################

#  save total score for summary paragraph
tot_tscore = df_ASRS.iloc[11, 1]
tot_percent = df_ASRS.iloc[11, 2]
tot_class = df_ASRS.iloc[11, 3]

##  delete total row (last row)
df_ASRS = df_ASRS.drop(index=11).reset_index(drop=True)

##  remove the space after / so that the 2 tables match
df_ASRS['Scale'] = df_ASRS['Scale'].str.replace('Social/ Communication', 'Social/Communication')

##  Add the description column to the table
ls_ASRSdescriptions = ['Indicates the extent to which the child uses verbal and nonverbal communication appropriately to initiate, engage in, and maintain, social contact.', 
                       'Indicates the child’s level of tolerance for changes in routine, engagement in apparently purposeless and stereotypical behaviors, and overreaction, to certain sensory experiences.', 
                       'Indicates deficits in attention and/or impulse control.',
                       'Indicates the child’s willingness and capacity to successfully engage in activities that develop and maintain relationships with other children.', 
                       'Indicates the child’s willingness and capacity to successfully engage in activities that develop and maintain relationships with adults.', 
                       'Indicates the child’s ability to provide an appropriate emotional response to another person in a social situation.', 
                       'Indicates the extent to which the child is able to utilize spoken communication in a structured and conventional way.', 
                       'Indicates the extent to which the child engages in apparently purposeless and repetitive behaviors.', 
                       'Indicates the extent to which the child tolerates changes in his environment, routines, activities, or behaviors.', 
                       'Indicates the child’s level of tolerance for certain experiences sensed through touch, sound, vision, smell, or taste.', 
                       'Indicates the extent to which the child is able to appropriately focus attention on one thing while ignoring other things, as well as how well the child controls his behavior and thoughts, maintains focus, and resists distraction.'
                      ]
df_ASRS['Descriptions'] = ls_ASRSdescriptions

##  function to combine strings
def combine_strings(row):
    return f"Ratings on this scale yielded a T-score of {row['T-score']} which is ranked at the {row['Percentile']} percentile and falls in the {row['Classification']} range."

# Apply the function to each row
df_ASRS['Scores'] = df_ASRS.apply(combine_strings, axis=1)

##  delete rows 1-4
df_ASRS = df_ASRS.drop(columns=df_ASRS.columns[[1,2,3,4]]).reset_index(drop=True)

# def add_comment(df_temp, col_name):
#     if df_temp['Item Score'].iloc[0] >= 3:
#         df_ASRS.loc[df_ASRS['Scale'] == col_name, 'Scores'] = df_ASRS.loc[df_ASRS['Scale'] == col_name, 'Scores'].values[0] + ' ' + df_temp[col_name].iloc[0]

def add_comment(df_temp, col_name):
    df_temp['Item Score'] = pd.to_numeric(df_temp['Item Score'], errors='coerce')
    if df_temp['Item Score'].iloc[0] >= 3:
        df_ASRS.loc[df_ASRS['Scale'] == col_name, 'Scores'] = df_ASRS.loc[df_ASRS['Scale'] == col_name, 'Scores'].values[0] + ' ' + df_temp[col_name].iloc[0]

add_comment(df_ub, 'Unusual Behaviors')
add_comment(df_sc, 'Social/Communication')
add_comment(df_sr, 'Self-Regulation')
add_comment(df_ps, 'Peer Socialization')
add_comment(df_as, 'Adult Socialization')
add_comment(df_ser, 'Social/Emotional Reciprocity')
add_comment(df_al, 'Atypical Language')
add_comment(df_ss, 'Sensory Sensitivity')
add_comment(df_s, 'Stereotypy')
add_comment(df_a, 'Attention')
add_comment(df_br, 'Behavioral Rigidity')


# group scales by classification (excluding Total Score)
df_filtered = df_ASRSappendix[df_ASRSappendix['Scale'] != 'Total Score']

very_elevated = df_filtered[df_filtered['Classification'] == 'Very Elevated Score']['Scale'].tolist()
elevated = df_filtered[df_filtered['Classification'] == 'Elevated Score']['Scale'].tolist()
slightly_elevated = df_filtered[df_filtered['Classification'] == 'Slightly Elevated Score']['Scale'].tolist()
average = df_filtered[df_filtered['Classification'] == 'Average Score']['Scale'].tolist()
low = df_filtered[df_filtered['Classification'] == 'Low']['Scale'].tolist()

def format_list(lst):
    if len(lst) == 0:
        return ''
    elif len(lst) == 1:
        return lst[0]
    else:
        return ', '.join(lst[:-1]) + ', and ' + lst[-1]

summary = f"Overall ASRS score: {firstName}'s total Score on the ASRS fell in the {tot_class} range, falling at the {tot_percent} percentile."

if very_elevated:
    summary += f" ASRS reflected very elevated scores in {format_list(very_elevated)}."
if elevated:
    summary += f" ASRS reflected elevated scores in {format_list(elevated)}."
if slightly_elevated:
    summary += f" Slightly elevated scores were reflected in {format_list(slightly_elevated)}."
if average:
    summary += f" An average score was reflected in {format_list(average)}."

doc = Document()

# Add table with header row + data rows
table = doc.add_table(rows=1, cols=len(df_ASRS.columns))
table.style = 'Table Grid'

# Set column widths (in inches converted to EMUs)
from docx.shared import Inches
table.columns[0].width = Inches(0.5)  # first column wide
table.columns[1].width = Inches(3.0)  # middle column narrow
table.columns[2].width = Inches(3.0)  # middle column narrow

# Write header row (bolded)
header_cells = table.rows[0].cells
for i, col_name in enumerate(df_ASRS.columns):
    header_cells[i].text = str(col_name)
    for run in header_cells[i].paragraphs[0].runs:
        run.bold = True
# Write data rows
for _, row in df_ASRS.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value) if value is not None else ''

# Add spacer and summary paragraph
doc.add_paragraph()
doc.add_paragraph(summary)

doc.save("C:/Users/ljkop/Laura/ASRS/ASRStableBody.docx")
